import os
from docx import Document
import pandas as pd

# Importa la libreria per l'API Gemini (sostituisci con il nome corretto se diverso)
import genai

# Importa pymilvus per gestire il database Milvus
from pymilvus import (
    connections,
    FieldSchema, CollectionSchema,
    DataType, Collection
)

##########################################
# 1. Lettura del file DOCX
##########################################
# Specifica il percorso del file DOCX
docx_file = "documento.docx"

# Leggi il file DOCX e unisci i paragrafi in un unico testo
document = Document(docx_file)
text = "\n".join([para.text for para in document.paragraphs if para.text.strip() != ""])

# Per esempio, possiamo usare il titolo del documento come primo paragrafo non vuoto
title = document.paragraphs[0].text if document.paragraphs else "Titolo non definito"

##########################################
# 2. Generazione degli embeddings con Gemini API
##########################################
# Imposta il modello da utilizzare (assicurati che sia corretto)
model = 'models/embedding-001'

# Genera gli embeddings: per documenti, task_type deve essere "retrieval_document".
# Nota: il parametro title migliora la qualità degli embeddings.
embedding_response = genai.embed_content(
    model=model,
    content=text,
    task_type="retrieval_document",
    title=title
)

# Estrai il vettore degli embeddings dalla risposta
embedding_vector = embedding_response["embedding"]
vector_dim = len(embedding_vector)
print("Embedding generato (dimensione {}):".format(vector_dim))
print(embedding_vector)

##########################################
# 3. Connessione a Milvus e creazione della collezione
##########################################
# Connetti a Milvus (modifica host e port se necessario)
connections.connect(host='localhost', port='19530')

# Definisci lo schema della collezione
id_field = FieldSchema(name="id", dtype=DataType.INT64, is_primary=True, auto_id=True)
title_field = FieldSchema(name="title", dtype=DataType.VARCHAR, max_length=200)
text_field = FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=5000)
embedding_field = FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=vector_dim)

schema = CollectionSchema(
    fields=[id_field, title_field, text_field, embedding_field],
    description="Collezione per memorizzare embeddings dei documenti DOCX"
)

# Crea o recupera la collezione
collection_name = "doc_embeddings"
try:
    collection = Collection(name=collection_name, schema=schema)
    print("Collezione '{}' creata.".format(collection_name))
except Exception as e:
    print("Errore nella creazione della collezione:", e)
    # Se la collezione esiste già, puoi recuperarla
    collection = Collection(name=collection_name)

##########################################
# 4. Inserimento dei dati nel database Milvus
##########################################
# Prepara i dati da inserire: organizziamo in un dizionario.
data = [
    [title],        # lista dei titoli
    [text],         # lista dei testi
    [embedding_vector]  # lista degli embeddings
]

# I dati devono essere ordinati seguendo l’ordine dei campi definiti nello schema:
# in questo caso: title, text, embedding (l'id verrà generato automaticamente)
# Se la collezione ha un campo id auto_id, non serve passarlo.
insert_result = collection.insert(data)
print("Inserimento completato. IDs inseriti:", insert_result.primary_keys)

##########################################
# 5. (Opzionale) Creazione di un indice per velocizzare le query
##########################################
index_params = {
    "index_type": "IVF_FLAT",
    "metric_type": "L2",
    "params": {"nlist": 128}
}
collection.create_index(field_name="embedding", index_params=index_params)
print("Indice creato sulla colonna 'embedding'.")

##########################################
# 6. (Opzionale) Salvataggio dello stato della collezione
##########################################
collection.load()
print("Collezione caricata e pronta per le query.")

# Ora hai letto il documento, generato gli embeddings e salvato i dati in Milvus.
